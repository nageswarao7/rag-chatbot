import uuid
import re
import numpy as np
import fitz  # PyMuPDF
from docx import Document
from io import BytesIO
from datetime import datetime
from sentence_transformers import SentenceTransformer
from backend.vector_store import vector_store

class SemanticChunker:
    """
    Advanced Semantic Chunking Strategy.
    
    Approach:
    ---------
    1. Split text into sentences
    2. Compute embeddings for each sentence
    3. Calculate similarity between adjacent sentences
    4. Detect topic boundaries where similarity drops
    5. Group sentences into semantically coherent chunks
    6. Merge small chunks, split large ones
    
    Why Semantic Chunking?
    ----------------------
    - Each chunk contains ONE coherent topic
    - Better retrieval accuracy (chunks match queries better)
    - No arbitrary mid-topic splits
    """
    
    def __init__(self):
        # Use same model as ChromaDB for consistency
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        self.similarity_threshold = 0.5  # Below this = topic change
        self.min_chunk_size = 300        # Minimum chars per chunk
        self.max_chunk_size = 3000       # Maximum chars per chunk (~750 tokens)
        self.target_chunk_size = 1500    # Ideal chunk size (~375 tokens)
    
    def split_into_sentences(self, text: str) -> list[str]:
        """Split text into sentences."""
        # Split on sentence boundaries
        sentences = re.split(r'(?<=[.!?])\s+', text)
        # Filter empty and very short sentences
        return [s.strip() for s in sentences if len(s.strip()) > 10]
    
    def compute_similarities(self, sentences: list[str]) -> list[float]:
        """Compute cosine similarity between adjacent sentences."""
        if len(sentences) < 2:
            return []
        
        # Get embeddings for all sentences
        embeddings = self.embedding_model.encode(sentences)
        
        # Calculate cosine similarity between adjacent sentences
        similarities = []
        for i in range(len(embeddings) - 1):
            sim = np.dot(embeddings[i], embeddings[i + 1]) / (
                np.linalg.norm(embeddings[i]) * np.linalg.norm(embeddings[i + 1])
            )
            similarities.append(float(sim))
        
        return similarities
    
    def find_breakpoints(self, similarities: list[float]) -> list[int]:
        """Find indices where topic changes (low similarity)."""
        breakpoints = []
        for i, sim in enumerate(similarities):
            if sim < self.similarity_threshold:
                breakpoints.append(i + 1)  # Break AFTER this sentence
        return breakpoints
    
    def create_chunks_from_breakpoints(self, sentences: list[str], breakpoints: list[int]) -> list[str]:
        """Group sentences into chunks based on breakpoints."""
        chunks = []
        start = 0
        
        for bp in breakpoints:
            chunk_text = ' '.join(sentences[start:bp])
            if chunk_text.strip():
                chunks.append(chunk_text)
            start = bp
        
        # Add remaining sentences
        if start < len(sentences):
            chunk_text = ' '.join(sentences[start:])
            if chunk_text.strip():
                chunks.append(chunk_text)
        
        return chunks
    
    def balance_chunks(self, chunks: list[str]) -> list[str]:
        """Merge small chunks and split large ones."""
        balanced = []
        buffer = ""
        
        for chunk in chunks:
            buffer += " " + chunk if buffer else chunk
            
            # If buffer exceeds max, split it
            if len(buffer) >= self.max_chunk_size:
                # Find a good split point
                split_point = buffer.rfind('. ', 0, self.max_chunk_size)
                if split_point == -1:
                    split_point = self.max_chunk_size
                
                balanced.append(buffer[:split_point + 1].strip())
                buffer = buffer[split_point + 1:].strip()
            
            # If buffer reaches target size and ends properly, flush it
            elif len(buffer) >= self.target_chunk_size and buffer.rstrip().endswith(('.', '!', '?')):
                balanced.append(buffer.strip())
                buffer = ""
        
        # Don't forget remaining buffer
        if buffer.strip():
            if balanced and len(buffer) < self.min_chunk_size:
                # Merge tiny remainder with previous chunk
                balanced[-1] += " " + buffer.strip()
            else:
                balanced.append(buffer.strip())
        
        return balanced
    
    def chunk(self, text: str) -> list[str]:
        """
        Main chunking method.
        Returns list of semantically coherent chunks.
        """
        if not text or len(text.strip()) < self.min_chunk_size:
            return [text.strip()] if text.strip() else []
        
        # Step 1: Split into sentences
        sentences = self.split_into_sentences(text)
        
        if len(sentences) <= 1:
            return [text.strip()]
        
        # Step 2: Compute similarities
        similarities = self.compute_similarities(sentences)
        
        # Step 3: Find topic breakpoints
        breakpoints = self.find_breakpoints(similarities)
        
        # Step 4: Create initial chunks
        chunks = self.create_chunks_from_breakpoints(sentences, breakpoints)
        
        # Step 5: Balance chunk sizes
        balanced_chunks = self.balance_chunks(chunks)
        
        return balanced_chunks


semantic_chunker = SemanticChunker()


class IngestionService:
    """Service for processing and ingesting documents with semantic chunking."""
    
    def __init__(self):
        self.chunker = semantic_chunker

    def extract_text_from_pdf(self, content: bytes) -> list[dict]:
        """Extract text from PDF with page metadata."""
        doc = fitz.open(stream=content, filetype="pdf")
        pages = []
        for page_num, page in enumerate(doc, 1):
            text = page.get_text()
            if text.strip():
                pages.append({
                    "text": text,
                    "section": f"Page {page_num}",
                    "page_number": page_num,
                    "total_pages": len(doc)
                })
        doc.close()
        return pages

    def extract_text_from_docx(self, content: bytes) -> list[dict]:
        """Extract text from DOCX with section metadata."""
        doc = Document(BytesIO(content))
        sections = []
        current = {"text": "", "section": "Introduction"}
        
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                if current["text"].strip():
                    sections.append(current)
                current = {"text": "", "section": para.text[:50]}
            current["text"] += para.text + "\n"
        
        if current["text"].strip():
            sections.append(current)
        
        return sections if sections else [{"text": "\n".join([p.text for p in doc.paragraphs]), "section": "Main"}]

    def extract_text_from_txt(self, content: bytes) -> list[dict]:
        """Extract text from plain text."""
        return [{"text": content.decode('utf-8'), "section": "Full Document"}]

    def process_document(self, content: bytes, filename: str) -> int:
        """
        Process document with semantic chunking.
        Returns number of chunks created.
        """
        ext = filename.lower().split('.')[-1]
        
        # Extract content
        if ext == 'pdf':
            sections = self.extract_text_from_pdf(content)
        elif ext in ['docx', 'doc']:
            sections = self.extract_text_from_docx(content)
        elif ext in ['txt', 'md', 'csv']:
            sections = self.extract_text_from_txt(content)
        else:
            raise ValueError(f"Unsupported: {ext}")
        
        if not sections:
            raise ValueError("No content extracted")
        
        all_chunks = []
        all_metas = []
        all_ids = []
        
        for section in sections:
            text = section.get("text", "")
            if not text.strip():
                continue
            
            # Apply SEMANTIC CHUNKING
            chunks = self.chunker.chunk(text)
            
            for idx, chunk_text in enumerate(chunks):
                chunk_id = str(uuid.uuid4())
                
                metadata = {
                    "source": filename,
                    "file_type": ext,
                    "ingested_at": datetime.now().isoformat(),
                    "section": section.get("section", ""),
                    "page_number": section.get("page_number", 0),
                    "chunk_index": idx,
                    "total_chunks": len(chunks),
                    "chunk_size": len(chunk_text),
                    "chunking_method": "semantic"
                }
                
                all_chunks.append(chunk_text)
                all_metas.append(metadata)
                all_ids.append(chunk_id)
        
        if all_chunks:
            vector_store.add_documents(documents=all_chunks, metadatas=all_metas, ids=all_ids)
        
        return len(all_chunks)


ingestion_service = IngestionService()
